import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
from fpdf import FPDF

class MailMergeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Mail Merge to PDF")
        self.docx_file = None
        self.xlsx_file = None
        self.create_widgets()
    
    def create_widgets(self):
        tk.Button(self.root, text="Select DOCX (DOTX) File", command=self.load_docx).pack(pady=5)
        tk.Button(self.root, text="Select XLSX File", command=self.load_xlsx).pack(pady=5)
        tk.Button(self.root, text="Perform Mail Merge", command=self.mail_merge).pack(pady=20)
    
    def load_docx(self):
        self.docx_file = filedialog.askopenfilename(filetypes=[("Word files", "*.docx *.dotx")])
        if self.docx_file:
            messagebox.showinfo("File Selected", f"DOCX/DOTX file selected: {self.docx_file}")
    
    def load_xlsx(self):
        self.xlsx_file = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if self.xlsx_file:
            self.df = pd.read_excel(self.xlsx_file)
            self.column = tk.simpledialog.askstring("Input", f"Available columns:\n{', '.join(self.df.columns)}\n\nEnter the column name for file names:")
            if self.column in self.df.columns:
                messagebox.showinfo("File Selected", f"XLSX file selected: {self.xlsx_file}\nColumn for file names: {self.column}")
            else:
                messagebox.showerror("Error", f"Column {self.column} not found in the Excel file.")
    
    def mail_merge(self):
        if not self.docx_file or not self.xlsx_file or not self.column:
            messagebox.showerror("Error", "Please select both DOCX/DOTX and XLSX files and specify the column for file names.")
            return
        
        doc_template = Document(self.docx_file)
        for _, row in self.df.iterrows():
            doc = Document()
            for element in doc_template.element.body:
                doc.element.body.append(element)
            
            for paragraph in doc.paragraphs:
                for key, value in row.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value))
            
            output_file = f"{row[self.column]}.pdf"
            self.save_as_pdf(doc, output_file)
            messagebox.showinfo("Success", f"Mail merge completed. Files saved with names based on {self.column} column.")
    
    def save_as_pdf(self, doc, filename):
        pdf = FPDF()
        pdf.add_page()
        for para in doc.paragraphs:
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, para.text)
        pdf.output(filename)

if __name__ == "__main__":
    root = tk.Tk()
    app = MailMergeApp(root)
    root.mainloop()